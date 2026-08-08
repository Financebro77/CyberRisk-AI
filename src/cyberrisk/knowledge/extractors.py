"""Format extractors — read a document file into clean text + sections.

Each extractor returns ``(clean_text, sections)`` where:

    clean_text  the full cleaned document text (paragraphs preserved),
    sections    a list of ``Section`` dataclass: {heading, text, start_char}
                describing the document's section structure.  ``start_char``
                is the char offset in ``clean_text``, so a chunk's section_ref
                can point at the exact location (honors chunk.schema.json).

Supported formats (stdlib for TXT/HTML/Markdown; optional libs for PDF/DOCX):

    txt    read text; first non-empty line is the implicit title
    html   strip tags via html.parser; headings become sections
    md     split on '# ' headings into sections (no dep)
    pdf    pypdf (optional 'knowledge' extra) — graceful error if absent
    docx   python-docx (optional 'knowledge' extra) — graceful error if absent

Every extractor cleans its output through ``cleaners.clean_text`` so all
downstream stages see consistent text.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
import re

from cyberrisk.knowledge.cleaners import clean_text


@dataclass(frozen=True)
class Section:
    """One heading + its following text block, with its char offset in the doc."""

    heading: str
    text: str
    start_char: int


@dataclass
class ExtractedDocument:
    """Result of extraction: clean text + section structure + implicit title."""

    text: str
    sections: list[Section] = field(default_factory=list)
    title: str = ""

    def char_span_for(self, section: Section) -> tuple[int, int]:
        """Character span [start, end) of a section's text within ``text``."""
        start = section.start_char
        return start, start + len(section.text)


def _implicit_title(text: str) -> str:
    """First non-empty line of the document, if it looks like a title."""
    for line in text.splitlines():
        line = line.strip()
        if line:
            # A single short line at the top is the title.  Clean it so smart
            # punctuation (em-dashes etc.) is folded to ASCII like the body.
            return clean_text(line)
    return ""


def extract_text(path: str | Path) -> str:
    """Read a file's raw text, given its extension.

    Raises
        RuntimeError  for PDF/DOCX when the required optional dependency is
                      not installed — a clear, actionable error (never a raw
                      ImportError).
        ValueError    for an unsupported extension.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix in (".md", ".markdown"):
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix in (".html", ".htm"):
        return _read_html(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(
        f"unsupported document extension {suffix!r} for {path.name}; "
        "supported: .txt, .md/.markdown, .html/.htm, .pdf, .docx"
    )


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def extract_markdown(path: str | Path) -> ExtractedDocument:
    """Extract markdown: headings become sections, body is the remaining text."""
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    text = clean_text(raw)
    return _sections_from_headings(text, _HEADING_RE)


# ---------------------------------------------------------------------------
# HTML (stdlib html.parser — no bs4 dependency)
# ---------------------------------------------------------------------------


class _TextHTMLParser(HTMLParser):
    """Collect text and heading tokens from an HTML document.

    ``<h1>..<h6>`` open a new section; text between headings belongs to the
    current section.  Script/style content is skipped.
    """

    _SKIP_TAGS = {"script", "style", "head", "title", "meta", "noscript"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[str, str]] = []  # (heading, text) in order
        self._current_heading: str | None = None
        self._current_text: list[str] = []
        self._skip_depth = 0
        self._in_heading = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
            return
        if re.fullmatch(r"h[1-6]", tag):
            self._flush_text()
            self._current_heading = ""
            self._in_heading = True

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._in_heading and re.fullmatch(r"h[1-6]", tag):
            self._in_heading = False
            self._flush_text()

    def handle_data(self, data: str) -> None:
        if self._skip_depth > 0:
            return
        if self._in_heading:
            self._current_heading = (self._current_heading or "") + data
        else:
            self._current_text.append(data)

    def _flush_text(self) -> None:
        heading = (self._current_heading or "").strip()
        text = "".join(self._current_text)
        if heading or text.strip():
            self.blocks.append((heading, text))
        self._current_heading = None
        self._current_text = []

    def finish(self) -> list[tuple[str, str]]:
        self._flush_text()
        return [(h, t) for h, t in self.blocks if h or t.strip()]


def extract_html(path: str | Path) -> ExtractedDocument:
    """Extract HTML: <h1..h6> headings become sections, body text is cleaned."""
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    parser = _TextHTMLParser()
    parser.feed(raw)
    parser.close()
    blocks = parser.finish()

    # Rebuild the clean full text from the blocks, tracking char offsets.
    parts: list[str] = []
    sections: list[Section] = []
    offset = 0
    for heading, body in blocks:
        block_text = clean_text(f"{heading}\n\n{body}" if heading else body)
        if not block_text:
            continue
        sections.append(Section(heading=heading, text=block_text, start_char=offset))
        parts.append(block_text)
        offset += len(block_text) + 2  # two-char paragraph separator
    text = clean_text("\n\n".join(parts))
    return ExtractedDocument(text=text, sections=sections, title=_implicit_title(text))


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def extract_txt(path: str | Path) -> ExtractedDocument:
    """Extract plain text: first non-empty line is the implicit title.

    No section structure (a plain document) — the 'plain' chunker handles it.
    """
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    text = clean_text(raw)
    return ExtractedDocument(text=text, sections=[], title=_implicit_title(text))


def extract_yaml_incident(path: str | Path) -> ExtractedDocument:
    """Extract a structured incident YAML as a narrative document.

    Loads the Incident (validated against the taxonomy) and renders its
    ``narrative()`` — a self-citing block with all ten fields — as the document
    text.  This lets an incident flow through the standard ingest/embed/RAG
    pipeline exactly like any other document, while the structured IncidentIndex
    still supports field-queried retrieval.
    """
    from cyberrisk.knowledge.incidents import load_incident

    incident = load_incident(path)
    text = clean_text(incident.narrative())
    return ExtractedDocument(
        text=text,
        sections=[],
        title=f"{incident.company} — {incident.attack_type}",
    )


# ---------------------------------------------------------------------------
# PDF (optional pypdf) and DOCX (optional python-docx)
# ---------------------------------------------------------------------------


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Reading PDFs requires 'pypdf'. Install the knowledge extra: "
            "`pip install -e '.[knowledge]'`."
        ) from exc
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "Reading DOCX requires 'python-docx'. Install the knowledge extra: "
            "`pip install -e '.[knowledge]'`."
        ) from exc
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_pdf(path: str | Path) -> ExtractedDocument:
    """Extract PDF text; no heading structure is recovered (pypdf gives pages)."""
    raw = _read_pdf(path)
    text = clean_text(raw)
    return ExtractedDocument(text=text, sections=[], title=_implicit_title(text))


def extract_docx(path: str | Path) -> ExtractedDocument:
    """Extract DOCX paragraphs; no heading structure is recovered."""
    raw = _read_docx(path)
    text = clean_text(raw)
    return ExtractedDocument(text=text, sections=[], title=_implicit_title(text))


# ---------------------------------------------------------------------------
# Heading-based section extraction (shared by markdown)
# ---------------------------------------------------------------------------


def _sections_from_headings(
    text: str,
    heading_re: re.Pattern[str],
) -> ExtractedDocument:
    """Split clean text on heading matches into sections.

    Each section = (heading, following body up to the next heading).  The
    section's ``start_char`` is the body's char offset in ``text``, so a
    chunk's char_span points at the body while section_ref carries the heading.
    The first heading is treated as the implicit title when it is the first
    line of the document.
    """
    matches = list(heading_re.finditer(text))
    if not matches:
        return ExtractedDocument(text=text, sections=[], title=_implicit_title(text))

    sections: list[Section] = []
    title = ""
    if text.splitlines() and matches[0].start() <= len(text.splitlines()[0]):
        title = matches[0].group(2).strip()

    for i, m in enumerate(matches):
        heading = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip("\n")
        sections.append(Section(heading=heading, text=body, start_char=start))

    return ExtractedDocument(text=text, sections=sections, title=title)
